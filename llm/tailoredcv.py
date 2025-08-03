

import time
import os
import json
from docx import Document
from docx.shared import Pt
from app.llm_utils import call_gpt
import re
from docx.shared import Pt
from difflib import get_close_matches

# === CONFIG ===
RANKED_DIR = "data/json/rank_candidates_by_job"
PARSED_DIR = "data/json/llm_normalized"
TEMPLATE_META_DIR = "templates/json_extracted_template"
OUTPUT_DIR = "data/json/llm_tailored_cv"
DOCX_OUTPUT_DIR = "data/json/tailored_cv_docx"
LLM_RESULTS_DIR = "data/json/llm_results"

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(DOCX_OUTPUT_DIR, exist_ok=True)

# === UTILS ===

def normalize_string(text):
    return re.sub(r'[^a-z0-9]', '', text.lower())

def load_template_structure():
    for file in os.listdir(TEMPLATE_META_DIR):
        if file.endswith("__extracted_template.json"):
            path = os.path.join(TEMPLATE_META_DIR, file)
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data.get("sections", {}), data.get("template_name")
    raise FileNotFoundError("❌ No extracted template JSON found.")

def find_resume_file(candidate_name):
    cand_normalized = normalize_string(candidate_name)
    for file in os.listdir(PARSED_DIR):
        fname_normalized = normalize_string(file.replace(".json", ""))
        if cand_normalized == fname_normalized:
            return os.path.normpath(os.path.join(PARSED_DIR, file))
    return None

def safe_call_gpt(system_prompt, user_prompt, retries=2):
    for attempt in range(retries):
        try:
            response = call_gpt(system_prompt, user_prompt)
            if response:
                return response.strip()
        except Exception as e:
            print(f"⚠️ GPT failed attempt {attempt+1}: {e}")
            time.sleep(2)
    return ""

def clean_heading_from_content(section_name, content):
    if not content:
        return ""
    pattern = rf"(?i)^(\s*(#+|\*+)?\s*{re.escape(section_name.strip())}[\s:/\-]*\n*)+"
    return re.sub(pattern, "", content.strip(), flags=re.IGNORECASE)

def get_matched_skills_from_llm_results(candidate_name, job_title):
    cand_normalized = normalize_string(candidate_name)
    matched_file = None
    for fname in os.listdir(LLM_RESULTS_DIR):
        if not fname.endswith(".json"):
            continue
        fname_base = fname.replace("__matches.json", "").replace(".json", "")
        fname_normalized = normalize_string(fname_base)
        if fname_normalized == cand_normalized:
            matched_file = os.path.join(LLM_RESULTS_DIR, fname)
            break

    matched_skills = []
    if matched_file and os.path.exists(matched_file):
        with open(matched_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        all_matches = data.get("matched_jobs", [])
        for match in all_matches:
            if normalize_string(match.get("job_role", "")) == normalize_string(job_title):
                must = match.get("matched", {}).get("must_have", {}).get("matched", [])
                nice = match.get("matched", {}).get("nice_to_have", {}).get("matched", [])
                bonus = match.get("matched", {}).get("bonus", {}).get("matched", [])
                matched_skills = must + nice + bonus
                break
    else:
        print(f"⚠️ Could not find matching LLM result file for: {candidate_name}")

    # Add GPT-suggested skills
    system_prompt = "You are an expert job analyst. Suggest only core skills."
    user_prompt = f"""
Suggest 6 to 7 important technical skills for a resume tailored to this job title:

Job Title: {job_title}

Return the skills as a comma-separated list only, no explanations.
"""
    gpt_response = safe_call_gpt(system_prompt, user_prompt)
    gpt_skills = [s.strip() for s in gpt_response.split(",") if s.strip()]

    all_skills = matched_skills + gpt_skills
    deduped = {}
    for skill in all_skills:
        lower = skill.lower()
        if lower not in deduped:
            deduped[lower] = skill
    return list(sorted(deduped.values()))

def get_certifications_from_parsed_resume(candidate_name):
    cand_normalized = normalize_string(candidate_name)
    for file in os.listdir(PARSED_DIR):
        if not file.endswith(".json"):
            continue
        fname_normalized = normalize_string(file.replace(".json", ""))
        if cand_normalized == fname_normalized:
            with open(os.path.join(PARSED_DIR, file), "r", encoding="utf-8") as f:
                data = json.load(f)
            certs = data.get("certifications", [])
            if isinstance(certs, list):
                return "\n- " + "\n- ".join(sorted(set(certs)))
            elif isinstance(certs, str):
                return certs.strip()
    return "[No certifications found]"




def gpt_fill_functional_skills(job_title):
    system_prompt = "You are a resume enhancement expert and job role analyst."
    user_prompt = f"""
The candidate is tailoring their resume for this job role: "{job_title}".

Based on this role, suggest 5 to 7 **functional skills** required to succeed in this job.

Return a Python list of short, role-relevant phrases. Keep it concise.
Examples: ["Agile", "Scrum", "Stakeholder Communication", "Problem Solving", "Client Management"]
Do not include technical tools or languages.
"""
    gpt_result = safe_call_gpt(system_prompt, user_prompt)
    try:
        return eval(gpt_result) if isinstance(eval(gpt_result), list) else [gpt_result.strip()]
    except:
        return [s.strip() for s in gpt_result.split(",") if s.strip()]



def gpt_fill_business_sector(job_title):
    system_prompt = "You are a job description analyst. Focus only on business industry domain."
    user_prompt = f"""
The candidate is applying for this job role: "{job_title}".

Based on the role and general industry trends, suggest 1 or 2 relevant **business sectors** or **industry domains** this job belongs to.

Examples: ["IT Services", "Telecommunications"], ["Banking"], ["Healthcare Technology"]

Return a clean Python list only. No explanations.
"""
    gpt_result = safe_call_gpt(system_prompt, user_prompt)
    try:
        return eval(gpt_result) if isinstance(eval(gpt_result), list) else [gpt_result.strip()]
    except:
        return [s.strip() for s in gpt_result.split(",") if s.strip()]


def get_experience_from_parsed_resume(candidate_name):
    cand_normalized = normalize_string(candidate_name)
    for file in os.listdir(PARSED_DIR):
        if not file.endswith(".json"):
            continue
        fname_normalized = normalize_string(file.replace(".json", ""))
        if cand_normalized == fname_normalized:
            with open(os.path.join(PARSED_DIR, file), "r", encoding="utf-8") as f:
                data = json.load(f)
            experience_entries = data.get("experience", [])
            if isinstance(experience_entries, list):
                formatted = ""
                for entry in experience_entries:
                    role = entry.get("role", "N/A")
                    company = entry.get("company", "N/A")
                    years = entry.get("years", "")
                    desc = entry.get("description", "").strip()
                    formatted += f"**{role}**, {company} ({years})\n- {desc}\n\n"
                return formatted.strip()
    return "[No experience found]"

def get_education_from_parsed_resume_or_gpt(candidate_name, job_title):
    cand_normalized = normalize_string(candidate_name)

    # === Try loading from llm_normalized/<candidate>.json ===
    for file in os.listdir(PARSED_DIR):
        if not file.endswith(".json"):
            continue
        fname_normalized = normalize_string(file.replace(".json", ""))
        if cand_normalized == fname_normalized:
            with open(os.path.join(PARSED_DIR, file), "r", encoding="utf-8") as f:
                data = json.load(f)
            education = data.get("education", [])
            if isinstance(education, list) and education:
                formatted = []
                for edu in education[:2]:  # ⛔ Limit to 2 entries
                    degree = edu.get("degree", "").strip()
                    institute = edu.get("institute", "").strip()
                    years = edu.get("years", "").strip()
                    if degree and institute:
                        formatted.append(f"**{degree}**, {institute} ({years})")
                if formatted:
                    return "\n\n".join(formatted)

    # === Fallback: GPT generates 2 relevant bachelor-level degrees ===
    system_prompt = "You are a resume generator filling the Education section."
    user_prompt = f"""
The candidate has no education listed. They're applying for the job role: "{job_title}".

Generate exactly **2 bachelor-level degrees**, relevant to the job. Do not include coursework or long paragraphs.

Format each like:
"**Bachelor of Science in XYZ**, University Name (Graduated: Year)"

Only return two entries. Keep it professional.
"""
    gpt_response = safe_call_gpt(system_prompt, user_prompt)
    if not gpt_response:
        return "[No education available]"

    # Cleanup: extract only 2 lines if GPT returns more
    lines = [line.strip() for line in gpt_response.splitlines() if line.strip()]
    return "\n\n".join(lines[:2]) if lines else "[No education available]"


def get_projects_from_parsed_resume_with_gpt_fill(candidate_name, job_title):
    cand_normalized = normalize_string(candidate_name)
    projects = []

    # Load candidate JSON
    for file in os.listdir(PARSED_DIR):
        if not file.endswith(".json"):
            continue
        fname_normalized = normalize_string(file.replace(".json", ""))
        if cand_normalized == fname_normalized:
            with open(os.path.join(PARSED_DIR, file), "r", encoding="utf-8") as f:
                data = json.load(f)
            projects = data.get("projects", [])
            break

    results = []
    for proj in projects[:3]:  # limit to top 3
        title = proj.get("title", "").strip()
        desc = proj.get("rojects", "").strip()

        if not title:
            continue

        # Rewrite weak/empty descriptions
        if not desc or len(desc) < 15:
            system_prompt = "You're a resume expert improving project relevance."
            user_prompt = f"""
Improve this resume project for job title "{job_title}".

Project Title: {title}
Current Description: "{desc or 'None'}"

Write a 1-2 line professional and job-relevant description.
Avoid fluff or tools unrelated to {job_title}.
"""
            desc = safe_call_gpt(system_prompt, user_prompt)

        results.append(f"**{title}**\n{desc.strip()}")

    return "\n\n".join(results) if results else "[No projects found]"


def evaluate_tailoring_score_with_gpt(tailored_data):
    system_prompt = "You are a professional resume evaluator and job relevance checker."
    user_prompt = f"""
The following resume has been tailored for the job role: "{tailored_data.get('TARGET_JOB', '')}".
Evaluate how well this tailored resume matches the job role in terms of relevance, content quality, and completeness.

Here is the tailored resume data (in JSON-style):
{json.dumps(tailored_data, indent=2)}

Rules:
- Check if sections are relevant to the job.
- Check if the content is professionally written and not vague or empty.
- Penalize for generic or irrelevant projects or skills.
- Reward for precise, relevant, well-filled sections.
- Do NOT include personal info (name, email, etc.) in judgment.

Only return a score from 1 to 10 as a single integer. No explanation.
"""
    score_str = safe_call_gpt(system_prompt, user_prompt)
    match = re.search(r"\b([1-9]|10)\b", score_str)
    return f"{match.group(1)} / 10" if match else "0 / 10"


def gpt_fill_section(section_name, resume_data, job_title, force_projects=False):
    system_prompt = "You are a professional resume writer."
    user_prompt = f"""
You are generating the section \"{section_name}\" for a resume tailored to the job role: \"{job_title}\".
Generate relevant, job-specific, and professional content based only on the candidate's background below. Avoid repetition and irrelevant domains.

Candidate Info:
Name: {resume_data.get("candidate_name")}
Experience: {resume_data.get("experience")}
skills: {resume_data.get("skills")}
Project Descriptions: {resume_data.get("projects")}
Certifications: {resume_data.get("certifications")}
"""
    if force_projects:
        user_prompt += "\n\nNOTE: Projects are missing or weak. Generate 3 strong, QA-relevant projects."

    debug_entry = {"section": section_name, "regenerated": False}
    response = safe_call_gpt(system_prompt, user_prompt)
    if not response or len(response.strip()) < 30 or response.lower() in ["n/a", "none", "null", "-"]:
        debug_entry["final"] = "❌ GPT failed to generate valid content."
        return "", debug_entry

    validator_system = "You are a resume quality inspector and job relevance checker."
    validator_prompt = f"""
SECTION: {section_name}
JOB: {job_title}
CONTENT:
{response}

RULES:
- Must match job domain
- No irrelevant tools
- No vague or placeholder text
Return only ✅ or ❌
"""
    verdict = safe_call_gpt(validator_system, validator_prompt)
    if "✅" in verdict:
        return response, {"section": section_name, "verdict": verdict}
    else:
        retry_prompt = user_prompt + "\n\nRephrase and ensure relevance."
        retry = safe_call_gpt(system_prompt, retry_prompt)
        return retry or "", {"section": section_name, "verdict": "❌", "regenerated": True}


def tailor_candidate_to_job(candidate_name, job_title, template_sections, template_name):
    print(f"\n🎯 Tailoring resume for: {candidate_name} → {job_title}")
    resume_path = find_resume_file(candidate_name)
    if not resume_path:
        print(f"❌ Resume file not found for {candidate_name}")
        return

    with open(resume_path, "r", encoding="utf-8") as f:
        resume_data = json.load(f)

    matched_skills = get_matched_skills_from_llm_results(candidate_name, job_title)
    if matched_skills:
        resume_data["skills"] = matched_skills

    tailored = {
        "CANDIDATE_NAME": candidate_name,
        "TARGET_JOB": job_title,
        "TEMPLATE_USED": template_name,
    }

    for section, required in template_sections.items():
        section_clean = section.strip()

        if section_clean == "Personal Info":
            pi = []
            for key in ["name", "email", "phone", "linkedin", "github"]:
                if resume_data.get(key):
                    pi.append(f"{key.capitalize()}: {resume_data[key]}")
            tailored[section_clean] = "\n".join(pi) if pi else "[Missing Personal Info]"
            continue
        if section_clean.lower() == "skills":
           skills_list = resume_data.get("skills", [])
           tailored[section_clean] = "- " + "\n- ".join(sorted(set(skills_list))) if skills_list else "[No matched skills found]"
           continue


        if section_clean == "Certifications":
            tailored[section_clean] = get_certifications_from_parsed_resume(candidate_name)
            continue

        if section_clean == "Functional Skills":
            functional_skills = gpt_fill_functional_skills(job_title)
            tailored[section_clean] = functional_skills
            continue

        if section_clean == "Business Sector":
            business_sector = gpt_fill_business_sector(job_title)
            tailored[section_clean] = business_sector
            continue

        if section_clean == "Languages":
            tailored[section_clean] = "- English\n- Urdu"
            continue

        if section_clean == "Experience":
            tailored[section_clean] = get_experience_from_parsed_resume(candidate_name)
            continue

        if section_clean == "Education":
            tailored[section_clean] = get_education_from_parsed_resume_or_gpt(candidate_name, job_title)
            continue

        if section_clean == "Project Descriptions":
            tailored[section_clean] = get_projects_from_parsed_resume_with_gpt_fill(candidate_name, job_title)
            continue

        # Fallback for other sections using GPT
        original = resume_data.get(section_clean)
        if original and isinstance(original, str) and len(original.strip()) > 30:
            tailored[section_clean] = clean_heading_from_content(section_clean, original)
        else:
            value, _ = gpt_fill_section(section_clean, resume_data, job_title)
            tailored[section_clean] = clean_heading_from_content(section_clean, value)



    # === Add Tailoring Score BEFORE saving
    tailored["TAILORING_SCORE"] = evaluate_tailoring_score_with_gpt(tailored)


    # === Final cleanup: keep only required keys
    template_keys = set(template_sections.keys()) | {
        "CANDIDATE_NAME", "TARGET_JOB", "TEMPLATE_USED", "GPT_FILLED", "TAILORING_SCORE"
    }
    tailored = {k: v for k, v in tailored.items() if k in template_keys}

    # === Save JSON
    json_filename = f"{candidate_name.lower().replace(' ', '_')}__{job_title.lower().replace(' ', '_')}__tailored.json"
    json_path = os.path.join(OUTPUT_DIR, json_filename)
    with open(json_path, "w", encoding="utf-8") as f:
      json.dump(tailored, f, indent=2)

    print(f"✅ Tailored resume saved: {json_path}")



# 🔁 Heading aliases: actual heading in Word : key from JSON
HEADING_ALIASES = {
    "profile": "Profile / Summary",
    "experience": "Work Experience",
    "education": "Education",
    "project description": "Projects",
    "certifications": "Certifications",
    "functional skills": "Functional Skills",
    "business sector": "Business Sector",
    "languages": "Languages"
}


def find_best_heading_match(doc, json_key):
    # Try to match based on alias
    for heading_text in HEADING_ALIASES:
        if HEADING_ALIASES[heading_text].lower() == json_key.lower():
            return heading_text

    # If no alias match, try fuzzy match
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") or p.text.strip().isupper()]
    matches = get_close_matches(json_key.strip(), headings, n=1, cutoff=0.5)
    return matches[0] if matches else None


def replace_docx_section_content(doc, json_key, new_content):
    match_heading = find_best_heading_match(doc, json_key)
    if not match_heading:
        print(f"⚠️ Could not find matching heading for: '{json_key}'")
        return

    inside = False
    i = 0

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if not inside and text.lower() == match_heading.lower():
            inside = True
            i += 1
            continue

        if inside:
            if para.style.name.startswith("Heading") or para.text.strip().isupper():
                break
            para._element.getparent().remove(para._element)
            doc.paragraphs.pop(i)
            continue

        i += 1

    for para in doc.paragraphs:
        if para.text.strip().lower() == match_heading.lower():
            for line in new_content.strip().split("\n"):
                if line.strip():
                    new_para = para.insert_paragraph_after(line.strip())
                    new_para.style.font.size = Pt(11)
            print(f"✅ Replaced section: '{json_key}' → heading: '{match_heading}'")
            break


def insert_personal_info_top(doc, info_string):
    first_para = doc.paragraphs[0]
    lines = info_string.strip().split("\n")
    for line in reversed(lines):
        para = first_para.insert_paragraph_before(line.strip())
        para.style.font.size = Pt(11)
    print("✅ Inserted Personal Info at the top.")



def create_new_docx_from_json_template(tailored_json_path, template_json_path, output_dir):
    if not os.path.exists(tailored_json_path):
        print(f"❌ Tailored JSON not found: {tailored_json_path}")
        return

    if not os.path.exists(template_json_path):
        print(f"❌ Template JSON not found: {template_json_path}")
        return

    with open(tailored_json_path, "r", encoding="utf-8") as f:
        tailored_data = json.load(f)

    with open(template_json_path, "r", encoding="utf-8") as f:
        template_data = json.load(f)

    doc = Document()

    # Add main heading
    doc.add_heading(f"Tailored Resume - {tailored_data.get('CANDIDATE_NAME', '')}", level=1)

    # Go section by section as per template JSON
    section_keys = template_data.get("sections", {})
    for section_name in section_keys:
        content = tailored_data.get(section_name, "").strip()
        if not content:
            continue
        doc.add_heading(section_name.strip(), level=2)
        for line in str(content).splitlines():
            if line.strip():
                doc.add_paragraph(line.strip(), style='Normal')

    # Save the file
    candidate_name = tailored_data.get("CANDIDATE_NAME", "unknown").lower().replace(" ", "_")
    job_title = tailored_data.get("TARGET_JOB", "job").lower().replace(" ", "_")
    output_path = os.path.join(output_dir, f"{candidate_name}__{job_title}__tailored.docx")
    doc.save(output_path)
    print(f"✅ DOCX created from scratch: {output_path}")



from docx import Document
from docx.shared import Pt
import os
import json

def write_tailored_docx_from_json(tailored_json_path, output_dir):
    if not os.path.exists(tailored_json_path):
        print(f"❌ JSON not found: {tailored_json_path}")
        return

    with open(tailored_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    candidate_name = data.get("CANDIDATE_NAME", "unknown").strip().replace(" ", "_").lower()
    job_title = data.get("TARGET_JOB", "job").strip().replace(" ", "_").lower()

    output_path = os.path.join(output_dir, f"{candidate_name}__{job_title}__tailored.docx")

    # Create a new Word document from scratch
    doc = Document()

    # Add title
    doc.add_heading(f"Tailored Resume for {data.get('CANDIDATE_NAME', '')}", level=1)

    # Write each meaningful section from tailored JSON
    for section, content in data.items():
        if section in ["GPT_FILLED", "TAILORING_SCORE", "CANDIDATE_NAME", "TARGET_JOB", "TEMPLATE_USED"]:
            continue  # Skip meta/meta fields

        doc.add_heading(section, level=2)
        for line in str(content).splitlines():
            doc.add_paragraph(line.strip(), style='Normal')

    doc.save(output_path)
    print(f"📄 Saved tailored DOCX: {output_path}")



def convert_tailored_json_to_docx(candidate_name, job_title):
    json_filename = f"{candidate_name.lower().replace(' ', '_')}__{job_title.lower().replace(' ', '_')}__tailored.json"
    tailored_json_path = os.path.join(OUTPUT_DIR, json_filename)

    write_tailored_docx_from_json(
        tailored_json_path=tailored_json_path,
        output_dir=DOCX_OUTPUT_DIR
    )



def tailor_all():
    template_sections, template_name = load_template_structure()
    docx_template_path = "templates/Word Template.docx"

    ranking_files = [f for f in os.listdir(RANKED_DIR) if f.endswith(".json")]
    if not ranking_files:
        print("❌ No ranking JSON files found in rank_candidates_by_job.")
        return

    for ranking_file in ranking_files:
        ranking_path = os.path.join(RANKED_DIR, ranking_file)
        print(f"\n📁 Processing ranking file: {ranking_file}")

        with open(ranking_path, "r", encoding="utf-8") as f:
            rankings = json.load(f)

        # Determine whether rankings is a list or dict
        if isinstance(rankings, list):
            for entry in rankings:
                job_title = entry.get("job_title", "")
                candidates = entry.get("ranked_candidates", [])
                if candidates and isinstance(candidates, list):
                    candidate_name = candidates[0].get("candidate_name", "")
                    if candidate_name:
                        tailor_candidate_to_job(candidate_name, job_title, template_sections, template_name)
                        convert_tailored_json_to_docx(candidate_name, job_title, )

        elif isinstance(rankings, dict):
            for job_title, candidates in rankings.items():
                if candidates and isinstance(candidates, list):
                    candidate_name = candidates[0].get("candidate_name", "")
                    if candidate_name:
                        tailor_candidate_to_job(candidate_name, job_title, template_sections, template_name)
                        convert_tailored_json_to_docx(candidate_name, job_title, )


if __name__ == "__main__":
    print("🔥 Starting Full CV Tailoring Process")
    tailor_all()



