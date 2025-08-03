import os
import json
import re
from docx import Document
from PyPDF2 import PdfReader
from app.llm_utils import call_gpt

TEMPLATE_DIR = "templates"
TEMPLATE_OUTPUT_DIR = "templates/json_extracted_template"
os.makedirs(TEMPLATE_OUTPUT_DIR, exist_ok=True)

def extract_docx_text(path):
    doc = Document(path)
    lines = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    return "\n".join(lines)

def extract_pdf_text(path):
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def parse_template_with_llm(template_name, raw_text):
    system_prompt = "You are an expert in resume template understanding and formatting."
    user_prompt = f"""
You are given the raw text of a resume template used by a company. Your job is to parse this into a clean structured JSON format, identifying sections such as:

- Personal Info (Name, Email, Address)
- Profile / Summary
- Skills (Technical, Platforms, Tools)
- Certifications
- Functional Skills
- Business Sector
- Languages
- Work Experience (Company, Title, Dates, Responsibilities)
- Education
- Projects

Output format:
{{
  "template_name": "{template_name}",
  "sections": {{
    ...
  }}
}}

Here is the raw template text:
-------------------------------
{raw_text}
"""

    response = call_gpt(system_prompt, user_prompt)

    # Strip markdown ```json wrapper if present
    match = re.search(r"```json\s*(\{.*?\})\s*```", response, re.DOTALL)
    cleaned = match.group(1) if match else response.strip()

    try:
        parsed = json.loads(cleaned)
    except Exception:
        print(f"⚠️ GPT returned malformed JSON. Wrapping raw response.")
        parsed = {
            "template_name": template_name,
            "sections": {
                "RAW": cleaned
            }
        }
    return parsed

def run_template_extractor():
    for file in os.listdir(TEMPLATE_DIR):
        if not file.lower().endswith((".docx", ".pdf")):
            continue

        template_name = file.rsplit(".", 1)[0]
        path = os.path.join(TEMPLATE_DIR, file)

        if file.lower().endswith(".docx"):
            print(f"📄 Extracting DOCX: {file}")
            raw_text = extract_docx_text(path)
        elif file.lower().endswith(".pdf"):
            print(f"📄 Extracting PDF: {file}")
            raw_text = extract_pdf_text(path)

        if not raw_text.strip():
            print(f"⚠️ No text extracted from {file}")
            continue

        parsed_json = parse_template_with_llm(template_name, raw_text)

        out_path = os.path.join(TEMPLATE_OUTPUT_DIR, f"{template_name}__extracted_template.json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(parsed_json, f, indent=2)

        print(f"✅ Parsed and saved: {out_path}")

if __name__ == "__main__":
    run_template_extractor()
